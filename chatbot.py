from typing import Dict, TypedDict, List
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage
from langgraph.graph import StateGraph, END
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_openai import ChatOpenAI
from langchain_community.tools import YouTubeSearchTool, ArxivQueryRun
from langchain_community.utilities import DuckDuckGoSearchAPIWrapper
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain.tools import Tool
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from dotenv import load_dotenv
import os
from langchain.agents import AgentExecutor, create_openai_functions_agent
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Load environment variables
load_dotenv()

# Initialize embeddings
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-mpnet-base-v2",
    model_kwargs={'device': 'cpu'},
    encode_kwargs={'normalize_embeddings': True}
)

# Initialize the model
model = ChatOpenAI(
    api_key=os.getenv("DEEP_SK_API_KEY"),
    base_url="https://api.deepseek.com/v1",
    model_name="deepseek-chat"
)

# Initialize tools with proper configurations
youtube_tool = YouTubeSearchTool(
    youtube_api_key=os.getenv("YOUTUBE_API_KEY"),
    number_results=3  # Limit to 3 results for cleaner output
)

arxiv_tool = ArxivQueryRun(
    top_k_results=3,
    load_max_docs=3,
    load_all_available_meta=True
)

search = DuckDuckGoSearchAPIWrapper(
    max_results=3
)

tavily_tool = TavilySearchResults(
    api_key=os.getenv("TAVILY_API_KEY"),
    max_results=3
)

# Create tool functions with better error handling and formatting
def youtube_search(query: str) -> str:
    """Search for YouTube videos and return formatted results."""
    try:
        results = youtube_tool.run(query)
        if not results:
            return "No videos found."
        return f"Here are some relevant videos:\n{results}"
    except Exception as e:
        return f"Error searching YouTube: {str(e)}"

def arxiv_search(query: str) -> str:
    """Search ArXiv and return formatted results."""
    try:
        results = arxiv_tool.run(query)
        if not results:
            return "No papers found."
        # Clean and format the results
        return f"Here are relevant papers from ArXiv:\n{results}"
    except Exception as e:
        return f"Error searching ArXiv: {str(e)}"

def web_search(query: str) -> str:
    """Search DuckDuckGo and return formatted results."""
    try:
        results = search.run(query)
        if not results:
            return "No results found."
        return f"Here's what I found:\n{results}"
    except Exception as e:
        return f"Error searching DuckDuckGo: {str(e)}"

def tavily_search(query: str) -> str:
    """Search Tavily and return formatted results."""
    try:
        results = tavily_tool.run(query)
        if not results:
            return "No detailed information found."
        return f"Here's the detailed information:\n{results}"
    except Exception as e:
        return f"Error searching Tavily: {str(e)}"

# Update tools list with new functions
tools = [
    Tool(
        name="YouTube Search",
        func=youtube_search,
        description="Search for YouTube videos. Use this when users want to find videos or tutorials."
    ),
    Tool(
        name="ArXiv Search",
        func=arxiv_search,
        description="Search academic papers on ArXiv. Use this for research papers and scientific publications."
    ),
    Tool(
        name="DuckDuckGo Search",
        func=web_search,
        description="Search the web using DuckDuckGo. Use this for general information and current events."
    ),
    Tool(
        name="Tavily Search",
        func=tavily_search,
        description="Get detailed analysis and information. Use this for comprehensive research and fact-checking."
    )
]

# Update agent prompt to be more specific about tool usage
agent_prompt = ChatPromptTemplate.from_messages([
    ("system", """You are a helpful AI assistant with access to multiple search tools.
    When using these tools:
    1. YouTube Search: Return actual video titles and links
    2. ArXiv Search: Return real paper titles, authors, and arxiv IDs
    3. DuckDuckGo Search: Return factual web search results
    4. Tavily Search: Return detailed analysis from verified sources
    
    Important:
    - Verify the results before presenting them
    - Format the results clearly and readably
    - If a tool returns no results or errors, try another relevant tool
    - Always provide context for the results you share"""),
    MessagesPlaceholder(variable_name="chat_history", optional=True),
    ("human", "{input}"),
    MessagesPlaceholder(variable_name="agent_scratchpad", optional=True)
]).partial()

# Create the agent
agent = create_openai_functions_agent(
    llm=model,
    prompt=agent_prompt,
    tools=tools
)

agent_executor = AgentExecutor.from_agent_and_tools(
    agent=agent,
    tools=tools,
    verbose=True,
    handle_parsing_errors=True
)

def process_pdf(pdf_file_path: str) -> FAISS:
    """Process PDF and create vector store"""
    loader = PyPDFLoader(pdf_file_path)
    documents = loader.load()
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    splits = text_splitter.split_documents(documents)
    
    vectorstore = FAISS.from_documents(splits, embeddings)
    return vectorstore


def chat_with_pdf(question: str, context: str, chat_history: List[BaseMessage] = None) -> str:
    """Generate response for PDF queries with memory"""
    if chat_history is None:
        chat_history = []
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a helpful AI assistant. Answer questions based on the provided context from the PDF document.
        If the answer cannot be found in the context, say so. Do not make up information.
        Use the chat history to provide more contextual and consistent responses."""),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", f"Context from PDF:\n{context}\n\nQuestion: {question}")
    ])
    
    messages = prompt.format_messages(chat_history=chat_history, context=context, question=question)
    response = model.invoke(messages)
    return response.content

def chat_with_memory(user_input: str, state: Dict = None):
    """Handle chat with memory using agent"""
    if state is None:
        state = {"messages": []}
    
    try:
        # Get response from agent
        response = agent_executor.invoke({
            "chat_history": state["messages"],
            "input": user_input
        })
        
        # Add messages to history
        state["messages"].append(HumanMessage(content=user_input))
        state["messages"].append(AIMessage(content=response["output"]))
        
        return response["output"], state
    except Exception as e:
        error_message = f"An error occurred: {str(e)}"
        state["messages"].append(HumanMessage(content=user_input))
        state["messages"].append(AIMessage(content=error_message))
        return error_message, state

def format_assignment_document(doc: Document, title: str, content: str) -> Document:
    """Format the assignment with proper styling"""
    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content with proper formatting
    content_sections = content.split('\n\n')
    for section in content_sections:
        para = doc.add_paragraph()
        para.add_run(section)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Inches(0.5)
        para.paragraph_format.line_spacing = 1.5
    
    return doc

def write_assignment(topic: str, requirements: str = None) -> tuple[str, Document]:
    """Generate formatted assignment content using local Blank.docx template"""
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are an expert academic writer and document formatter. Create a well-structured academic assignment.

        Follow these guidelines implicitly without adding any formatting tags or symbols:
        1. Begin with:
           [The provided name and registration number, or use placeholder if not provided]
           Topic/Subject
           (Leave a blank line after this section)
        
        2. Write a clear title centered on its own line
        
        3. Structure your content with:
           - Numbered sections (1., 2., 3.)
           - Each section title on its own line
           - Clear paragraph breaks between sections
           - Proper transitions between ideas
        
        4. Include:
           - Introduction with context and objectives
           - Well-developed main sections
           - Clear conclusion
           - References in APA format
        
        Write naturally and let the document structure flow through proper spacing and numbering.
        Do not include any formatting tags, asterisks, or special symbols."""),
        ("human", f"Topic: {topic}\nStudent Details: {requirements if requirements else 'Use placeholder student details'}")
    ])
    
    messages = prompt.format_messages(topic=topic, requirements=requirements)
    response = model.invoke(messages)
    
    # Use local Blank.docx
    doc = Document("Blank.docx")
    
    # Clear template content
    for paragraph in doc.paragraphs[:]:
        p = paragraph._element
        p.getparent().remove(p)
    
    # Process content with automatic formatting
    content = response.content
    lines = content.split('\n')
    
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        
        # Format student details and title
        if i <= 2:  # First three lines (name, regno, topic)
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
        # Rest of the formatting remains the same
        elif line.strip().startswith(('1.', '2.', '3.', '4.', '5.')):
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
        else:
            run = p.add_run(line)
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.5
    
    return response.content, doc 